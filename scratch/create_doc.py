import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_physical_data_model_doc():
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("3.1.3 Mô hình dữ liệu mức vật lý")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = 'Times New Roman'
    p_title.paragraph_format.space_after = Pt(14)
    
    tables_data = [
        {
            "name": "Loại thực thể Users",
            "desc": "Lưu thông tin tài khoản người dùng trong hệ thống (Khách hàng, Nhân viên, Quản trị viên), hỗ trợ đăng nhập OAuth2 Google và xác thực OTP.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã người dùng (Primary Key)"],
                ["name", "varchar(100)", "", "", "X", "Họ tên người dùng"],
                ["email", "varchar(100)", "", "X", "X", "Email đăng nhập"],
                ["password", "varchar(255)", "", "", "X", "Mật khẩu đã mã hóa (Bcrypt)"],
                ["phone", "varchar(20)", "", "", "", "Số điện thoại liên hệ"],
                ["avatar", "varchar(255)", "", "", "", "Đường dẫn ảnh đại diện"],
                ["points", "int", "", "", "X", "Điểm thưởng tích lũy"],
                ["membership_level", "enum", "", "", "X", "Phân hạng thành viên (new, silver, gold, diamond)"],
                ["role", "enum", "", "", "X", "Vai trò (customer, staff, admin)"],
                ["staff_type", "enum", "", "", "", "Phân loại nhân viên (reception, delivery)"],
                ["is_active", "tinyint(1)", "", "", "X", "Trạng thái tài khoản (1: Mở, 0: Khóa)"],
                ["google_id", "varchar(255)", "", "", "", "Mã xác thực tài khoản Google OAuth2"],
                ["remember_token", "varchar(100)", "", "", "", "Token ghi nhớ phiên đăng nhập"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo tài khoản"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật thông tin"]
            ]
        },
        {
            "name": "Loại thực thể UserAddresses",
            "desc": "Lưu danh sách sổ địa chỉ nhận hàng của người dùng kèm tọa độ vĩ độ/kinh độ trên bản đồ.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã địa chỉ"],
                ["user_id", "int", "", "", "X", "Mã người dùng (Khóa ngoại trỏ users)"],
                ["recipient_name", "varchar(100)", "", "", "X", "Tên người nhận hàng"],
                ["phone", "varchar(20)", "", "", "X", "Số điện thoại người nhận"],
                ["address_line", "text", "", "", "X", "Địa chỉ giao hàng chi tiết"],
                ["latitude", "decimal(10,8)", "", "", "", "Vĩ độ định vị địa lý trên bản đồ"],
                ["longitude", "decimal(11,8)", "", "", "", "Kinh độ định vị địa lý trên bản đồ"],
                ["is_default", "tinyint(1)", "", "", "X", "Đánh dấu địa chỉ mặc định (1: Đúng, 0: Sai)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo địa chỉ"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể Categories",
            "desc": "Quản lý danh mục nhóm đồ uống trong hệ thống.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã danh mục"],
                ["name", "varchar(100)", "", "", "X", "Tên danh mục đồ uống"],
                ["image_url", "varchar(255)", "", "", "", "Đường dẫn hình ảnh đại diện danh mục"],
                ["is_active", "tinyint(1)", "", "", "X", "Trạng thái hiển thị (1: Mở, 0: Ẩn)"],
                ["display_order", "int", "", "", "X", "Thứ tự ưu tiên hiển thị"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo danh mục"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể Products",
            "desc": "Lưu thông tin chi tiết từng sản phẩm đồ uống bán trên hệ thống.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã sản phẩm"],
                ["sku", "varchar(50)", "", "X", "X", "Mã quản lý kho sản phẩm (SKU)"],
                ["slug", "varchar(255)", "", "X", "X", "Đường dẫn xem chi tiết sản phẩm"],
                ["name", "varchar(255)", "", "", "X", "Tên đồ uống"],
                ["base_price", "decimal(10,2)", "", "", "X", "Giá bán gốc chưa tính size/topping"],
                ["image", "varchar(255)", "", "", "", "Đường dẫn hình ảnh đại diện"],
                ["description", "text", "", "", "", "Bài viết mô tả chi tiết đồ uống"],
                ["is_active", "tinyint(1)", "", "", "X", "Trạng thái kinh doanh (1: Mở bán, 0: Ngừng)"],
                ["category_id", "int", "", "", "X", "Mã danh mục (Khóa ngoại trỏ categories)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo sản phẩm"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể ProductImages",
            "desc": "Lưu album tập hợp nhiều hình ảnh phụ của từng sản phẩm.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã hình ảnh album"],
                ["product_id", "int", "", "", "X", "Mã sản phẩm (Khóa ngoại trỏ products)"],
                ["image_path", "varchar(255)", "", "", "X", "Đường dẫn file hình ảnh phụ"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tải lên"]
            ]
        },
        {
            "name": "Loại thực thể ProductSizes",
            "desc": "Quản lý danh mục kích cỡ (Size S, M, L) và chênh lệch giá theo size của đồ uống.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã kích cỡ sản phẩm"],
                ["product_id", "int", "", "", "X", "Mã sản phẩm (Khóa ngoại trỏ products)"],
                ["size_name", "varchar(50)", "", "", "X", "Tên kích cỡ (S, M, L)"],
                ["price_adjustment", "decimal(10,2)", "", "", "X", "Số tiền cộng thêm so với giá gốc"]
            ]
        },
        {
            "name": "Loại thực thể Toppings",
            "desc": "Danh mục các món topping ăn kèm đồ uống.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã topping"],
                ["name", "varchar(255)", "", "", "X", "Tên món topping (Trân châu, Thạch, Pudding)"],
                ["price", "decimal(10,2)", "", "", "X", "Giá bán topping"],
                ["is_available", "tinyint(1)", "", "", "X", "Trạng thái khả dụng (1: Còn, 0: Hết)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể ProductToppings",
            "desc": "Bảng liên kết trung gian xác định danh sách topping được phép áp dụng cho từng sản phẩm.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["product_id", "int", "X", "", "X", "Mã sản phẩm (Khóa ngoại trỏ products)"],
                ["topping_id", "int", "X", "", "X", "Mã topping (Khóa ngoại trỏ toppings)"]
            ]
        },
        {
            "name": "Loại thực thể Favorites",
            "desc": "Lưu danh sách các sản phẩm yêu thích (thả tim) của khách hàng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã lượt yêu thích"],
                ["user_id", "int", "", "", "X", "Mã người dùng (Khóa ngoại trỏ users)"],
                ["product_id", "int", "", "", "X", "Mã sản phẩm (Khóa ngoại trỏ products)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian thêm yêu thích"]
            ]
        },
        {
            "name": "Loại thực thể Carts",
            "desc": "Quản lý giỏ hàng mua sắm của từng khách hàng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã giỏ hàng"],
                ["user_id", "int", "", "", "", "Mã người dùng (nếu đã đăng nhập)"],
                ["session_id", "varchar(100)", "", "", "", "Mã Session (nếu khách vãng lai)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian khởi tạo giỏ hàng"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể CartItems",
            "desc": "Chi tiết các món đồ uống nằm trong giỏ hàng kèm chọn lựa đường, đá, size.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã dòng sản phẩm trong giỏ"],
                ["cart_id", "int", "", "", "X", "Mã giỏ hàng (Khóa ngoại trỏ carts)"],
                ["product_id", "int", "", "", "X", "Mã sản phẩm (Khóa ngoại trỏ products)"],
                ["size_name", "varchar(50)", "", "", "X", "Size đồ uống đã chọn (S, M, L)"],
                ["quantity", "int", "", "", "X", "Số lượng ly đặt mua"],
                ["sugar_level", "varchar(20)", "", "", "", "Tỷ lệ đường đã chọn (0%, 50%, 100%)"],
                ["ice_level", "varchar(20)", "", "", "", "Tỷ lệ đá đã chọn (Không đá, Nửa đá, Đủ đá)"],
                ["unit_price", "decimal(10,2)", "", "", "X", "Đơn giá món tại thời điểm thêm vào giỏ"],
                ["created_at", "timestamp", "", "", "X", "Thời gian thêm món vào giỏ"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể CartItemToppings",
            "desc": "Lưu thông tin chi tiết các topping kèm theo cho từng ly đồ uống trong giỏ hàng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã dòng topping trong giỏ"],
                ["cart_item_id", "int", "", "", "X", "Mã chi tiết giỏ hàng (Khóa ngoại cart_items)"],
                ["topping_id", "int", "", "", "X", "Mã topping (Khóa ngoại trỏ toppings)"],
                ["price", "decimal(10,2)", "", "", "X", "Đơn giá topping"]
            ]
        },
        {
            "name": "Loại thực thể Orders",
            "desc": "Quản lý thông tin tổng quan các đơn hàng đặt mua đồ uống trên hệ thống.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã định danh đơn hàng"],
                ["order_code", "varchar(50)", "", "X", "X", "Mã đơn hàng hiển thị (VD: HPY-881901)"],
                ["user_id", "int", "", "", "", "Mã người dùng đặt hàng (Khóa ngoại users)"],
                ["delivery_staff_id", "int", "", "", "", "Mã nhân viên giao hàng được phân công"],
                ["customer_name", "varchar(100)", "", "", "X", "Tên người nhận hàng"],
                ["customer_phone", "varchar(20)", "", "", "X", "Số điện thoại nhận hàng"],
                ["delivery_address", "text", "", "", "", "Địa chỉ giao hàng chi tiết"],
                ["delivery_latitude", "decimal(10,7)", "", "", "", "Vĩ độ địa chỉ giao hàng"],
                ["delivery_longitude", "decimal(10,7)", "", "", "", "Kinh độ địa chỉ giao hàng"],
                ["total_amount", "decimal(10,2)", "", "", "X", "Tổng giá trị đồ uống gốc"],
                ["discount_amount", "decimal(10,2)", "", "", "X", "Số tiền được giảm giá (Coupon)"],
                ["shipping_fee", "decimal(10,2)", "", "", "X", "Phí vận chuyển tính theo khoảng cách (km)"],
                ["weather_fee", "decimal(10,2)", "", "", "X", "Phụ thu thời tiết xấu (Mưa to, giông bão)"],
                ["final_amount", "decimal(10,2)", "", "", "X", "Tổng số tiền thanh toán thực tế"],
                ["payment_status", "enum", "", "", "X", "Trạng thái thanh toán (unpaid, paid)"],
                ["payment_method", "varchar(20)", "", "", "X", "Phương thức thanh toán (cod, momo)"],
                ["status", "enum", "", "", "X", "Trạng thái đơn (pending, confirmed, shipping, completed, cancelled)"],
                ["cancel_reason", "varchar(500)", "", "", "", "Lý do hủy đơn hàng"],
                ["coupon_code", "varchar(50)", "", "", "", "Mã giảm giá đã áp dụng"],
                ["promotion_id", "int", "", "", "", "Mã khuyến mãi (Khóa ngoại promotions)"],
                ["delivery_type", "enum", "", "", "X", "Hình thức nhận (delivery, pickup)"],
                ["distance_km", "decimal(5,2)", "", "", "", "Khoảng cách giao hàng tính bằng Kilomet"],
                ["customer_note", "text", "", "", "", "Ghi chú thêm của khách khi đặt món"],
                ["created_at", "timestamp", "", "", "X", "Thời gian đặt hàng"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật đơn hàng"]
            ]
        },
        {
            "name": "Loại thực thể OrderItems",
            "desc": "Chi tiết các ly đồ uống được mua trong từng đơn hàng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã chi tiết món trong đơn hàng"],
                ["order_id", "int", "", "", "X", "Mã đơn hàng (Khóa ngoại trỏ orders)"],
                ["product_id", "int", "", "", "X", "Mã sản phẩm (Khóa ngoại trỏ products)"],
                ["size_name", "varchar(50)", "", "", "X", "Size đồ uống (S, M, L)"],
                ["quantity", "int", "", "", "X", "Số lượng ly mua"],
                ["unit_price", "decimal(10,2)", "", "", "X", "Giá bán món tại thời điểm đặt"],
                ["sugar_level", "varchar(20)", "", "", "", "Mức đường chọn mua"],
                ["ice_level", "varchar(20)", "", "", "", "Mức đá chọn mua"]
            ]
        },
        {
            "name": "Loại thực thể OrderItemToppings",
            "desc": "Chi tiết các món topping mua đi kèm từng ly đồ uống trong đơn hàng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã chi tiết topping trong đơn"],
                ["order_item_id", "int", "", "", "X", "Mã chi tiết đơn (Khóa ngoại order_items)"],
                ["topping_id", "int", "", "", "X", "Mã topping (Khóa ngoại trỏ toppings)"],
                ["topping_name", "varchar(255)", "", "", "X", "Tên topping lưu lại lịch sử"],
                ["topping_price", "decimal(10,2)", "", "", "X", "Đơn giá topping tại thời điểm mua"]
            ]
        },
        {
            "name": "Loại thực thể Promotions",
            "desc": "Quản lý các chương trình ưu đãi và mã giảm giá (Coupon).",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã khuyến mãi"],
                ["code", "varchar(50)", "", "X", "X", "Mã nhập giảm giá (VD: NEWBIE)"],
                ["type", "enum", "", "", "X", "Loại giảm giá (fixed: cố định, percent: %)"],
                ["value", "decimal(10,2)", "", "", "X", "Giá trị số tiền giảm hoặc % giảm"],
                ["min_order_amount", "decimal(10,2)", "", "", "", "Giá trị đơn tối thiểu để áp dụng"],
                ["max_discount_amount", "decimal(10,2)", "", "", "", "Giảm giá tối đa (dành cho loại %)"],
                ["start_at", "datetime", "", "", "", "Thời gian bắt đầu mã ưu đãi"],
                ["end_at", "datetime", "", "", "", "Thời gian hết hạn mã ưu đãi"],
                ["is_active", "tinyint(1)", "", "", "X", "Trạng thái kích hoạt (1: Bật, 0: Tắt)"],
                ["usage_limit", "int", "", "", "", "Giới hạn số lượt dùng tối đa"],
                ["used_count", "int", "", "", "X", "Số lượt khách hàng đã sử dụng"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo mã"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể Reviews",
            "desc": "Quản lý các bài đánh giá, số sao và nhận xét đồ uống từ người mua.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã bài đánh giá"],
                ["user_id", "int", "", "", "X", "Mã khách hàng đánh giá (Khóa ngoại users)"],
                ["product_id", "int", "", "", "X", "Mã sản phẩm được đánh giá (Khóa ngoại products)"],
                ["order_id", "int", "", "", "", "Mã đơn hàng tương ứng (Khóa ngoại orders)"],
                ["rating", "tinyint", "", "", "X", "Số sao đánh giá (1 đến 5 sao)"],
                ["comment", "text", "", "", "", "Nội dung bài nhận xét"],
                ["image_url", "varchar(255)", "", "", "", "Ảnh chụp thực tế sản phẩm đính kèm"],
                ["is_visible", "tinyint(1)", "", "", "X", "Trạng thái duyệt hiển thị (1: Hiện, 0: Ẩn)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian viết đánh giá"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể Materials",
            "desc": "Quản lý danh mục các loại nguyên liệu pha chế trong kho của cửa hàng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "bigint", "X", "X", "X", "Mã nguyên liệu kho"],
                ["name", "varchar(191)", "", "", "X", "Tên nguyên liệu (Sữa tươi, Trà đen, Ly,...)"],
                ["unit", "varchar(191)", "", "", "X", "Đơn vị tính (Chai, Hộp, Bao 1kg, Lốc)"],
                ["unit_price", "decimal(12,2)", "", "", "X", "Đơn giá bình quân nhập vào"],
                ["current_stock", "decimal(10,2)", "", "", "X", "Số lượng tồn kho hiện tại"],
                ["is_active", "tinyint(1)", "", "", "X", "Trạng thái nguyên liệu (1: Dùng, 0: Ngừng)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian khởi tạo"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật tồn"]
            ]
        },
        {
            "name": "Loại thực thể MaterialImports",
            "desc": "Quản lý lịch sử nhập kho các lô nguyên liệu kèm thời hạn sử dụng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "bigint", "X", "X", "X", "Mã lô nhập kho"],
                ["material_id", "bigint", "", "", "X", "Mã nguyên liệu (Khóa ngoại materials)"],
                ["quantity", "decimal(10,2)", "", "", "X", "Số lượng nhập ban đầu của lô"],
                ["remaining_quantity", "decimal(10,2)", "", "", "", "Số lượng còn lại trong lô"],
                ["total_price", "decimal(12,2)", "", "", "X", "Tổng chi phí nhập lô nguyên liệu"],
                ["expiration_date", "date", "", "", "", "Ngày hết hạn sử dụng của lô nguyên liệu"],
                ["note", "varchar(191)", "", "", "", "Ghi chú lô nhập"],
                ["created_at", "timestamp", "", "", "X", "Thời gian nhập kho"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể InventoryMovements",
            "desc": "Nhật ký theo dõi lịch sử biến động xuất/nhập/hủy nguyên liệu trong kho.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "bigint", "X", "X", "X", "Mã nhật ký biến động kho"],
                ["material_id", "bigint", "", "", "X", "Mã nguyên liệu (Khóa ngoại materials)"],
                ["material_import_id", "bigint", "", "", "", "Mã lô nhập tương ứng (Khóa ngoại material_imports)"],
                ["order_id", "int", "", "", "", "Mã đơn hàng tiêu thụ nguyên liệu (Khóa ngoại orders)"],
                ["type", "enum", "", "", "X", "Loại biến động (import, order_reserve, dispose, adjustment)"],
                ["quantity", "decimal(12,3)", "", "", "X", "Số lượng tăng (+) hoặc giảm (-)"],
                ["unit_cost", "decimal(12,2)", "", "", "X", "Đơn giá chi phí nguyên liệu"],
                ["note", "varchar(191)", "", "", "", "Ghi chú biến động (VD: Hàng hết hạn)"],
                ["created_at", "timestamp", "", "", "X", "Thời điểm phát sinh biến động"]
            ]
        },
        {
            "name": "Loại thực thể Banners",
            "desc": "Quản lý banner bài viết quảng cáo hiển thị slider trên trang chủ.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã banner"],
                ["title", "varchar(255)", "", "", "X", "Tiêu đề chính của banner"],
                ["title_tag", "varchar(191)", "", "", "", "Nhãn phụ hiển thị trên banner"],
                ["image_url", "varchar(255)", "", "", "X", "Đường dẫn file ảnh banner trên máy tính"],
                ["mobile_image_url", "varchar(255)", "", "", "", "Đường dẫn ảnh banner phiên bản điện thoại"],
                ["link_url", "varchar(255)", "", "", "", "Đường dẫn liên kết khi bấm vào banner"],
                ["position", "varchar(50)", "", "", "X", "Vị trí hiển thị (VD: home_slider)"],
                ["display_order", "int", "", "", "X", "Thứ tự ưu tiên xuất hiện"],
                ["is_active", "tinyint(1)", "", "", "X", "Trạng thái hiển thị (1: Bật, 0: Tắt)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể Settings",
            "desc": "Cấu hình chung cho toàn bộ hệ thống (Khoảng cách, phí ship/km, phụ thu thời tiết, giờ mở/đóng cửa).",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã cấu hình"],
                ["key", "varchar(100)", "", "X", "X", "Khóa cấu hình (VD: store_latitude, weather_surcharge_enabled)"],
                ["value", "text", "", "", "", "Giá trị cấu hình lưu giữ"],
                ["type", "varchar(50)", "", "", "", "Kiểu dữ liệu của cấu hình (string, boolean, decimal)"],
                ["description", "text", "", "", "", "Ghi chú công dụng của tham số cấu hình"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        }
    ]
    
    for item in tables_data:
        # Heading
        h = doc.add_paragraph()
        r_h = h.add_run(item["name"])
        r_h.bold = True
        r_h.font.size = Pt(13)
        r_h.font.name = 'Times New Roman'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        
        # Desc
        p_desc = doc.add_paragraph()
        r_desc = p_desc.add_run("Mô tả: " + item["desc"])
        r_desc.italic = True
        r_desc.font.size = Pt(11.5)
        r_desc.font.name = 'Times New Roman'
        p_desc.paragraph_format.space_after = Pt(6)
        
        # Table
        table = doc.add_table(rows=1, cols=len(item["cols"]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Style table borders
        tblPr = table._tbl.tblPr
        borders = parse_xml(r'''
            <w:tblBorders %s>
                <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, title in enumerate(item["cols"]):
            hdr_cells[i].text = title
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [2, 3, 4] else (WD_ALIGN_PARAGRAPH.LEFT)
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            # Shading header background
            shading = parse_xml(r'<w:shd %s w:fill="E8EEF5"/>' % nsdecls('w'))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
            
        # Rows
        for r_idx, row_data in enumerate(item["rows"]):
            row_cells = table.add_row().cells
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = val
                p = row_cells[c_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [2, 3, 4] else (WD_ALIGN_PARAGRAPH.LEFT)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.5)
                    if c_idx in [2, 3, 4] and val == 'X':
                        run.font.bold = True
                        
        # Column Widths
        col_widths = [Inches(1.3), Inches(1.1), Inches(0.4), Inches(0.4), Inches(0.4), Inches(2.9)]
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
                
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    output_path = r"c:\Users\ADMIN\Desktop\DA_LVTN\LUAN-VAN-TOT-NGHIEP\3.1.3_Mo_hinh_du_lieu_muc_vat_ly.docx"
    doc.save(output_path)
    print("SAVED_SUCCESSFULLY:", output_path)

if __name__ == "__main__":
    create_physical_data_model_doc()
