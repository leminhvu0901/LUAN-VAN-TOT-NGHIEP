import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_missing_entities_doc():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Các loại thực thể còn thiếu trong Mục 3.1.3")
    run_title.bold = True
    run_title.font.size = Pt(15)
    p_title.paragraph_format.space_after = Pt(12)
    
    tables_data = [
        {
            "name": "Loại thực thể UserAddresses",
            "desc": "Lưu danh sách sổ địa chỉ nhận hàng của người dùng kèm tọa độ vĩ độ/kinh độ trên bản đồ.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã địa chỉ"],
                ["user_id", "int", "", "", "X", "Mã người dùng (Khóa ngoại trỏ users)"],
                ["recipient_name", "varchar(100)", "", "", "X", "Tên người nhận hàng"],
                ["phone", "varchar(20)", "", "", "X", "Số điện thoại người nhận"],
                ["address_line", "text", "", "", "X", "Địa chỉ giao hàng chi tiết"],
                ["latitude", "decimal(10,8)", "", "", "", "Vĩ độ định vị địa lý trên bản đồ"],
                ["longitude", "decimal(11,8)", "", "", "", "Kinh độ định vị địa lý trên bản đồ"],
                ["is_default", "tinyint(1)", "", "", "X", "Đánh dấu địa chỉ mặc định (1: Đúng, 0: Sai)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo địa chỉ"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể ProductImages",
            "desc": "Lưu album tập hợp nhiều hình ảnh phụ của từng sản phẩm.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã hình ảnh album"],
                ["product_id", "int", "", "", "X", "Mã sản phẩm (Khóa ngoại trỏ products)"],
                ["image_path", "varchar(255)", "", "", "X", "Đường dẫn file hình ảnh phụ"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tải lên"]
            ]
        },
        {
            "name": "Loại thực thể Materials",
            "desc": "Quản lý danh mục các loại nguyên liệu pha chế trong kho của cửa hàng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "bigint", "X", "X", "X", "Mã nguyên liệu kho"],
                ["name", "varchar(191)", "", "", "X", "Tên nguyên liệu (Sữa tươi, Trà đen, Ly,...)"],
                ["unit", "varchar(191)", "", "", "X", "Đơn vị tính (Chai, Hộp, Bao 1kg, Lốc)"],
                ["unit_price", "decimal(12,2)", "", "", "X", "Đơn giá bình quân nhập vào"],
                ["current_stock", "decimal(10,2)", "", "", "X", "Số lượng tồn kho hiện tại"],
                ["is_active", "tinyint(1)", "", "", "X", "Trạng thái nguyên liệu (1: Dùng, 0: Ngừng)"],
                ["created_at", "timestamp", "", "", "X", "Thời gian khởi tạo"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật tồn"]
            ]
        },
        {
            "name": "Loại thực thể MaterialImports",
            "desc": "Quản lý lịch sử nhập kho các lô nguyên liệu kèm thời hạn sử dụng.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "bigint", "X", "X", "X", "Mã lô nhập kho"],
                ["material_id", "bigint", "", "", "X", "Mã nguyên liệu (Khóa ngoại materials)"],
                ["quantity", "decimal(10,2)", "", "", "X", "Số lượng nhập ban đầu của lô"],
                ["remaining_quantity", "decimal(10,2)", "", "", "", "Số lượng còn lại trong lô"],
                ["total_price", "decimal(12,2)", "", "", "X", "Tổng chi phí nhập lô nguyên liệu"],
                ["expiration_date", "date", "", "", "", "Ngày hết hạn sử dụng của lô nguyên liệu"],
                ["note", "varchar(191)", "", "", "", "Ghi chú lô nhập"],
                ["created_at", "timestamp", "", "", "X", "Thời gian nhập kho"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        },
        {
            "name": "Loại thực thể InventoryMovements",
            "desc": "Nhật ký theo dõi lịch sử biến động xuất/nhập/hủy nguyên liệu trong kho.",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "bigint", "X", "X", "X", "Mã nhật ký biến động kho"],
                ["material_id", "bigint", "", "", "X", "Mã nguyên liệu (Khóa ngoại materials)"],
                ["material_import_id", "bigint", "", "", "", "Mã lô nhập tương ứng (Khóa ngoại material_imports)"],
                ["order_id", "int", "", "", "", "Mã đơn hàng tiêu thụ nguyên liệu (Khóa ngoại orders)"],
                ["type", "enum", "", "", "X", "Loại biến động (import, order_reserve, dispose, adjustment)"],
                ["quantity", "decimal(12,3)", "", "", "X", "Số lượng tăng (+) hoặc giảm (-)"],
                ["unit_cost", "decimal(12,2)", "", "", "X", "Đơn giá chi phí nguyên liệu"],
                ["note", "varchar(191)", "", "", "", "Ghi chú biến động (VD: Hàng hết hạn)"],
                ["created_at", "timestamp", "", "", "X", "Thời điểm phát sinh biến động"]
            ]
        },
        {
            "name": "Loại thực thể Settings",
            "desc": "Cấu hình chung cho toàn bộ hệ thống (Khoảng cách, phí ship/km, phụ thu thời tiết, giờ mở/đóng cửa).",
            "cols": ["Thuộc tính", "Kiểu dữ liệu", "K", "U", "M", "Diễn giải"],
            "rows": [
                ["id", "int", "X", "X", "X", "Mã cấu hình"],
                ["key", "varchar(100)", "", "X", "X", "Khóa cấu hình (VD: store_latitude, weather_surcharge_enabled)"],
                ["value", "text", "", "", "", "Giá trị cấu hình lưu giữ"],
                ["type", "varchar(50)", "", "", "", "Kiểu dữ liệu của cấu hình (string, boolean, decimal)"],
                ["description", "text", "", "", "", "Ghi chú công dụng của tham số cấu hình"],
                ["created_at", "timestamp", "", "", "X", "Thời gian tạo"],
                ["updated_at", "timestamp", "", "", "X", "Thời gian cập nhật"]
            ]
        }
    ]
    
    for item in tables_data:
        h = doc.add_paragraph()
        r_h = h.add_run(item["name"])
        r_h.bold = True
        r_h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        
        p_desc = doc.add_paragraph()
        r_desc = p_desc.add_run("Mô tả: " + item["desc"])
        r_desc.italic = True
        r_desc.font.size = Pt(11.5)
        p_desc.paragraph_format.space_after = Pt(6)
        
        table = doc.add_table(rows=1, cols=len(item["cols"]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        tblPr = table._tbl.tblPr
        borders = parse_xml(r'''
            <w:tblBorders %s>
                <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)
        
        hdr_cells = table.rows[0].cells
        for i, title in enumerate(item["cols"]):
            hdr_cells[i].text = title
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(11)
            shading = parse_xml(r'<w:shd %s w:fill="E8EEF5"/>' % nsdecls('w'))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
            
        for r_idx, row_data in enumerate(item["rows"]):
            row_cells = table.add_row().cells
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = val
                p = row_cells[c_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    if c_idx in [2, 3, 4] and val == 'X':
                        run.font.bold = True
                        
        col_widths = [Inches(1.3), Inches(1.1), Inches(0.4), Inches(0.4), Inches(0.4), Inches(2.9)]
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
                
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    output_path = r"c:\Users\ADMIN\Desktop\DA_LVTN\LUAN-VAN-TOT-NGHIEP\3.1.3_Cac_thuc_the_con_thieu.docx"
    doc.save(output_path)
    print("SAVED_MISSING_SUCCESSFULLY:", output_path)

if __name__ == "__main__":
    create_missing_entities_doc()
